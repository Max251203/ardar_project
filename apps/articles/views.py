# apps/articles/views.py
from django.http import HttpResponseForbidden
from apps.comments.forms import CommentForm
from apps.articles.models import Article
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, JsonResponse, HttpResponseForbidden
from django.shortcuts import get_object_or_404
import re
import io
import os
import requests
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.core.paginator import Paginator
from django.template.loader import render_to_string
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from .forms import ArticleForm
from .models import Article
from apps.comments.models import Comment, Rating
from apps.comments.forms import CommentForm, RatingForm
from django.db.models import Avg
from gtts import gTTS
from docx import Document
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage


def article_list(request):
    """
    Список всех статей с поиском, фильтрацией по дате и языку, и пагинацией
    """
    search_query = request.GET.get('q', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')
    language = request.GET.get('language', '')
    page = request.GET.get('page', 1)

    # Базовый запрос
    if request.user.is_authenticated:
        articles = (Article.objects.filter(is_approved=True) |
                    Article.objects.filter(author=request.user)).distinct()
    else:
        articles = Article.objects.filter(is_approved=True)

    # Применяем фильтры
    if search_query:
        articles = articles.filter(title__icontains=search_query)

    if date_from:
        articles = articles.filter(created_at__gte=date_from)

    if date_to:
        articles = articles.filter(created_at__lte=date_to)

    # Фильтр по языку
    if language:
        articles = articles.filter(language=language)

    # Сортировка
    articles = articles.order_by('-created_at')

    # Пагинация
    paginator = Paginator(articles, 6)  # 6 статей на страницу

    try:
        articles = paginator.page(page)
    except PageNotAnInteger:
        articles = paginator.page(1)
    except EmptyPage:
        articles = paginator.page(paginator.num_pages)

    return render(request, 'articles/list.html', {
        'articles': articles,
        'search_query': search_query,
        'date_from': date_from,
        'date_to': date_to,
        'language': language
    })


def article_detail(request, pk):
    """
    Детальный просмотр статьи с комментариями и рейтингом
    """
    article = get_object_or_404(Article, pk=pk)

    # Проверка доступа к неопубликованной статье
    if not article.is_approved:
        if not request.user.is_authenticated or (
            request.user != article.author and
            getattr(request.user, 'role', '') != 'admin' and
            not request.user.is_superuser
        ):
            return render(request, 'core/403.html', status=403)

    # Получаем оценку текущего пользователя
    user_rating = None
    if request.user.is_authenticated:
        user_rating = Rating.objects.filter(
            article=article, user=request.user).first()

    # Средний рейтинг
    avg_rating = Rating.objects.filter(
        article=article).aggregate(Avg('value'))['value__avg']

    # Все оценки
    ratings = Rating.objects.filter(article=article).select_related('user')

    # Комментарии
    comments = Comment.objects.filter(article=article).order_by('-created_at')

    # Один комментарий от текущего пользователя (если есть)
    user_comment = None
    if request.user.is_authenticated:
        user_comment = Comment.objects.filter(
            article=article, author=request.user).first()

    # Обработка формы комментария
    comment_form = None
    if request.user.is_authenticated:
        if request.method == 'POST':
            if user_comment:
                # Обновление существующего комментария
                comment_form = CommentForm(request.POST, instance=user_comment)
            else:
                # Добавление нового комментария
                comment_form = CommentForm(request.POST)

            if comment_form.is_valid():
                comment = comment_form.save(commit=False)
                comment.article = article
                comment.author = request.user
                comment.save()
                return redirect('article_detail', pk=article.pk)
        else:
            # Предзаполнение формы
            if user_comment:
                comment_form = CommentForm(instance=user_comment)
            else:
                comment_form = CommentForm()

    return render(request, 'articles/detail.html', {
        'article': article,
        'user_rating': user_rating,
        'avg_rating': avg_rating,
        'ratings': ratings,
        'comments': comments,
        'user_comment': user_comment,
        'comment_form': comment_form
    })


@login_required
def article_create(request):
    """
    Создание новой статьи:
    - поддержка ввода вручную
    - загрузка файлов (.txt, .docx, .pdf)
    - автоматическое одобрение привилегированным и админам
    """
    if not (getattr(request.user, 'role', '') in ['registered', 'privileged', 'admin'] or request.user.is_superuser):
        return render(request, 'core/403.html', status=403)

    form = ArticleForm(request.POST or None, request.FILES or None)

    if request.method == 'POST' and form.is_valid():
        article = form.save(commit=False)
        article.author = request.user

        # Обработка загруженного изображения
        if 'image' in request.FILES:
            image_file = request.FILES['image']
            # Читаем данные изображения
            article.image_data = image_file.read()
            article.image_name = image_file.name
            article.image_type = image_file.content_type

        if getattr(request.user, 'role', '') in ['privileged', 'admin'] or request.user.is_superuser:
            article.is_approved = True

        article.save()

        # Обработка файла (если есть)
        if 'uploaded_file' in request.FILES:
            extract_text_from_file(article, request.FILES['uploaded_file'])

        return redirect('article_detail', pk=article.pk)

    return render(request, 'admin_panel/article_form.html', {'form': form})


def detect_language(text):
    """
    Определяет язык текста на основе анализа символов.
    Поддерживаемые языки: русский (ru), английский (en), армянский (hy).
    Возвращает код языка или None, если язык не определен.
    """
    import re

    # Регулярные выражения для разных языков
    armenian_regex = re.compile(r'[\u0530-\u058F]')
    russian_regex = re.compile(r'[\u0400-\u04FF]')
    english_regex = re.compile(r'[a-zA-Z]')
    other_script_regex = re.compile(
        r'[^\u0530-\u058F\u0400-\u04FFa-zA-Z0-9\s.,!?;:\'"()\-]')

    # Подсчитываем количество символов каждого языка
    armenian_count = len(armenian_regex.findall(text))
    russian_count = len(russian_regex.findall(text))
    english_count = len(english_regex.findall(text))
    other_script_count = len(other_script_regex.findall(text))

    # Общее количество символов (без пробелов и знаков препинания)
    total_chars = armenian_count + russian_count + english_count + other_script_count

    if total_chars == 0:
        return None

    # Если есть символы других письменностей (китайский, арабский и т.д.)
    if other_script_count > 0 and other_script_count / total_chars > 0.1:
        # Если более 10% символов - другие письменности, считаем язык неподдерживаемым
        return None

    # Определяем преобладающий язык (должен составлять не менее 60% текста)
    if armenian_count > 0 and armenian_count / total_chars >= 0.6:
        return 'hy'
    elif russian_count > 0 and russian_count / total_chars >= 0.6:
        return 'ru'
    elif english_count > 0 and english_count / total_chars >= 0.6:
        return 'en'
    else:
        # Если нет явного преобладания, возвращаем None
        return None


def extract_text_from_file(article, uploaded_file):
    """
    Извлекает текст из файла (txt, docx) и сохраняет его в поле `text`.
    """
    if not uploaded_file:
        return

    ext = os.path.splitext(uploaded_file.name)[1].lower()
    text = ""

    try:
        if ext == '.txt':
            text = uploaded_file.read().decode('utf-8')
        elif ext == '.docx':
            # Сохраняем файл временно
            temp_path = os.path.join('media', 'temp', uploaded_file.name)
            os.makedirs(os.path.dirname(temp_path), exist_ok=True)

            with open(temp_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)

            # Извлекаем текст
            doc = Document(temp_path)
            text = '\n'.join([p.text for p in doc.paragraphs])

            # Удаляем временный файл
            os.remove(temp_path)
        else:
            print("[!] Нераспознаваемый тип файла:", ext)
            return
    except Exception as e:
        print(f"[!] Ошибка при извлечении текста из файла: {e}")
        return

    if text:
        # Преобразуем текст в HTML-формат для CKEditor
        # Разбиваем на абзацы и оборачиваем каждый в <p>
        paragraphs = text.split('\n\n')
        html_paragraphs = []

        for p in paragraphs:
            if p.strip():  # Пропускаем пустые абзацы
                html_p = p.replace('\n', '<br>')
                html_paragraphs.append(f'<p>{html_p}</p>')

        html_text = ''.join(html_paragraphs)
        article.text = html_text
        article.save()


@csrf_exempt
def process_file(request):
    """
    Обрабатывает загруженный файл и возвращает извлеченный текст.
    Поддерживаются только .txt и .docx форматы.
    """
    if request.method != 'POST' or 'file' not in request.FILES:
        return JsonResponse({'error': 'No file provided'}, status=400)

    uploaded_file = request.FILES['file']
    ext = os.path.splitext(uploaded_file.name)[1].lower()
    text = ""

    try:
        if ext == '.txt':
            text = uploaded_file.read().decode('utf-8')
        # apps/articles/views.py (продолжение)
        elif ext == '.docx':
            # Сохраняем файл временно
            temp_path = os.path.join('media', 'temp', uploaded_file.name)
            os.makedirs(os.path.dirname(temp_path), exist_ok=True)

            with open(temp_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)

            # Извлекаем текст
            doc = Document(temp_path)
            text = '\n'.join([p.text for p in doc.paragraphs])

            # Удаляем временный файл
            os.remove(temp_path)
        else:
            return JsonResponse({'error': 'Unsupported file format. Only .txt and .docx are supported.'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

    # Преобразуем текст в HTML
    paragraphs = text.split('\n\n')
    html_paragraphs = []

    for p in paragraphs:
        if p.strip():  # Пропускаем пустые абзацы
            html_p = p.replace('\n', '<br>')
            html_paragraphs.append(f'<p>{html_p}</p>')

    html_text = ''.join(html_paragraphs)

    return JsonResponse({'text': html_text})


# В apps/articles/views.py обновим функцию generate_audio_armtts

@login_required
def generate_audio_armtts(request, pk):
    """
    Генерирует аудио-версию статьи с использованием ArmTTS через RapidAPI
    или альтернативные методы
    """
    article = get_object_or_404(Article, pk=pk)

    if not article.text:
        return HttpResponse("Нет текста для озвучивания", status=400)

    # Очищаем HTML-теги для получения чистого текста
    clean_text = re.sub(r'<.*?>', ' ', article.text)
    clean_text = re.sub(r'\s+', ' ', clean_text).strip()

    # Ограничиваем длину текста, если он слишком большой
    if len(clean_text) > 500:
        clean_text = clean_text[:500]

    # Проверяем язык статьи
    article_language = article.language

    # Если язык не армянский, используем gTTS вместо ArmTTS API
    if article_language != 'hy':
        try:
            # Используем Google Text-to-Speech
            language_code = {
                'ru': 'ru',
                'en': 'en',
                'hy': 'hy'  # Может не поддерживаться
            }.get(article_language, 'en')

            tts = gTTS(text=clean_text, lang=language_code, slow=False)

            # Сохраняем аудио во временный файл
            audio_file = io.BytesIO()
            tts.write_to_fp(audio_file)
            audio_file.seek(0)

            # Возвращаем аудио
            return HttpResponse(audio_file.read(), content_type='audio/mpeg')

        except Exception as e:
            return HttpResponse(f"Ошибка при генерации аудио: {str(e)}", status=500)

    # Получаем API ключ из настроек
    from apps.core.models import SiteSettings
    api_key = SiteSettings.get_setting(
        'ARMTTS_API_KEY', '1e8d32c7c3msh767635ff925bcd7p13000fjsn07bbb0ccda3f')

    # Пробуем разные endpoints для ArmTTS API
    endpoints_to_try = [
        "https://armtts1.p.rapidapi.com/synthesize",
        "https://armtts1.p.rapidapi.com/v1/synthesize",
        "https://armtts1.p.rapidapi.com/v2/synthesize",
        "https://armtts1.p.rapidapi.com/v3/synthesize"
    ]

    # Пробуем ArmTTS API
    for url in endpoints_to_try:
        try:
            payload = {
                "text": clean_text,
                "voice": "male"
            }

            headers = {
                "content-type": "application/json",
                "X-RapidAPI-Key": api_key,
                "X-RapidAPI-Host": "armtts1.p.rapidapi.com"
            }

            response = requests.post(
                url, json=payload, headers=headers, timeout=30)

            if response.status_code == 200:
                return HttpResponse(response.content, content_type='audio/mpeg')
            elif response.status_code == 404:
                continue
            else:
                return HttpResponse(f"Ошибка API: {response.status_code} - {response.text}", status=500)

        except Exception as e:
            continue

    # Если все endpoints не работают, возвращаем ошибку
    return HttpResponse("ArmTTS API временно недоступен. Попробуйте использовать браузерную озвучку.", status=503)


def get_article_image(request, pk):
    """
    Возвращает изображение статьи из БД
    """
    article = get_object_or_404(Article, pk=pk)

    if not article.image_data:
        return HttpResponse("Изображение не найдено", status=404)

    return HttpResponse(article.image_data, content_type=article.image_type)


@login_required
def remove_article_image(request, pk):
    """
    Удаляет изображение из статьи
    """
    article = get_object_or_404(Article, pk=pk)

    # Проверка прав доступа
    if not (request.user == article.author or
            getattr(request.user, 'role', '') == 'admin' or
            request.user.is_superuser):
        return HttpResponseForbidden()

    # Удаляем данные изображения
    article.image_data = None
    article.image_name = None
    article.image_type = None
    article.save()

    return JsonResponse({'success': True})
